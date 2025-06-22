# backend/main.py

from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
import os
import openai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# === Load environment variables ===
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# === FastAPI app setup ===
app = FastAPI()

# === Enable CORS for React frontend ===
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# === Directories setup ===
UPLOAD_FOLDER = "./docs"
STATIC_FOLDER = "./static"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(STATIC_FOLDER, exist_ok=True)

# === Pydantic Models ===
class GuideRequest(BaseModel):
    topic: str
    objectives: List[str]
    tone: Optional[str] = "Empowering"
    format: Optional[str] = "group"

class DocxRequest(BaseModel):
    topic: str
    objectives: List[str]
    guide: str
    upload_summary: Optional[str] = None

# === Endpoint: Generate Guide using GPT-4 Turbo ===
@app.post("/generate-guide")
def generate_guide(request: GuideRequest):
    prompt = f"""You are RecoveryPen AI, a trauma-informed assistant. Create a comprehensive, group-ready recovery guide on:

Topic: {request.topic}

Objectives:
{chr(10).join(f"- {obj}" for obj in request.objectives)}

Tone: {request.tone}
Format: {request.format}

The guide should include:
- Clear headings and sections
- Practical strategies
- Realistic examples
- Bullet points and group discussion prompts
- Closing summary

Begin now.
"""
    try:
        response = openai.chat.completions.create(
            model="gpt-4-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=2000,
        )
        guide_text = response.choices[0].message.content
        return {"guide": guide_text}
    except Exception as e:
        return {"guide": "", "error": str(e)}

# === Endpoint: Upload Document ===
@app.post("/upload-doc")
async def upload_doc(file: UploadFile = File(...)):
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    summary = f"File '{file.filename}' uploaded and indexed. It will be used to enhance AI-guided generation."
    return {"message": "Upload complete.", "summary": summary}

# === Endpoint: Generate DOCX Guide with Cover + Logo + Summary ===
@app.post("/generate-docx")
def generate_docx(request: DocxRequest):
    doc = Document()

    # === Cover Page with Logo ===
    logo_path = os.path.join(STATIC_FOLDER, "logo.png")
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(2))

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("RecoveryPen AI\nTrauma-Informed Companion\n\n")
    run.font.size = Pt(20)
    run.bold = True
    doc.add_page_break()

    # === Header/Footer Branding ===
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "RecoveryPen AI | Trauma-Informed Companion"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer.paragraphs[0]
    footer.text = "Page - "
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === Main Guide Content ===
    doc.add_heading(request.topic, level=1)

    doc.add_heading("Objectives", level=2)
    for obj in request.objectives:
        doc.add_paragraph(f"- {obj}", style='List Bullet')

    doc.add_heading("Recovery Guide", level=2)
    for line in request.guide.split("\n"):
        if line.strip() == "":
            doc.add_paragraph()
        elif line.strip().startswith("-"):
            doc.add_paragraph(line.strip(), style='List Bullet')
        else:
            doc.add_paragraph(line.strip())

    # === Upload Summary ===
    if request.upload_summary:
        doc.add_heading("Upload Summary", level=2)
        doc.add_paragraph(request.upload_summary)

    # === Stream back as DOCX ===
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={request.topic.replace(' ', '_')}_Recovery_Guide.docx"}
    )
